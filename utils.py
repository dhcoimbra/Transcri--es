import os
import pandas as pd
import time
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from assemblyai import upload_audio, transcrever_audio_assemblyai
from PIL import UnidentifiedImageError, Image
from db import salvar_transcricao, buscar_transcricao
import streamlit as st
from docx.enum.text import WD_ALIGN_PARAGRAPH
import hashlib
import re

def redimensionar_imagem(image_path, max_width=7):
    """Redimensiona a imagem para ter uma largura máxima de 7 cm, mantendo a proporção."""
    with Image.open(image_path) as img:
        width, height = img.size
        aspect_ratio = height / width
        new_width = min(width, max_width * 37.7953)  # Conversão para pixels (1 cm ≈ 37.7953 pixels)
        new_height = int(new_width * aspect_ratio)
        img = img.resize((int(new_width), new_height), Image.LANCZOS)
        # Salvar a imagem temporária redimensionada
        temp_path = f"{image_path}_temp_resized.jpg"
        img.save(temp_path, format="JPEG", quality=85)
        return temp_path

# Função para adicionar bordas a todas as células da tabela
def adicionar_bordas_a_tabela(table):
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')  # Define as bordas como simples
                border.set(qn('w:sz'), '4')        # Espessura da borda
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')  # Cor preta
                tc_borders.append(border)
            
            tc_pr.append(tc_borders)

def verificar_arquivos_na_pasta(file, audio_dir):
    # Carregar o arquivo Excel
    df = pd.read_excel(file, engine='openpyxl', header=1)

    # Filtrar linhas em branco
    df = df.dropna(how='all')

    # Iterar pelas linhas e verificar se pelo menos um arquivo existe na pasta
    for index, row in df.iterrows():
        attachment = row.get('Attachment #1', None)
        if pd.notna(attachment):
            file_path = os.path.join(audio_dir, attachment)
            if os.path.exists(file_path):
                return True  # Arquivo encontrado, a pasta está correta
    return False  # Nenhum arquivo encontrado, a pasta está incorreta

def gerar_hash_arquivo(caminho_arquivo):
    hash_sha256 = hashlib.sha256()
    with open(caminho_arquivo, "rb") as f:
        # Lê o arquivo em blocos para calcular o hash (eficiente para arquivos grandes)
        for bloco in iter(lambda: f.read(4096), b""):
            hash_sha256.update(bloco)
    return hash_sha256.hexdigest()

def criar_documento_para_lote(df_lote, audio_dir, lote_num, progress_bar, linhas_processadas, total_linhas):
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'From'
    hdr_cells[2].text = 'Body'
    hdr_cells[3].text = 'Timestamp-Time'
    item = 0
    # Iterar por todas as linhas do DataFrame original
           
    for index, row in df_lote.iterrows():
        item = item + 1
        from_field = row['From']
        to_field = row.get('To', '')
        timestamp = row.get('Timestamp-Time', row.get('Timestamp: Time', ''))
        #timestamp = row.get('Timestamp-Time', '')
        body = row.get('Body', '')
        attachment = row.get('Attachment #1', None)
        #forwarded = row.get('Label', None)
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(item)
        # Preencher a coluna "From"
        row_cells[1].text = from_field

        if pd.notna(attachment):
            file_path = os.path.join(audio_dir, attachment)
            
            if attachment.endswith('.opus'):
                # Verifica se o arquivo existe
                if os.path.exists(file_path):
                    # Gerar hash do arquivo se ele existir
                    hash_arquivo = gerar_hash_arquivo(file_path)
                    transcricao_existente = buscar_transcricao(hash_arquivo)
                    
                    if transcricao_existente:
                        row_cells[2].text = f"ÁUDIO\nTranscrição: {transcricao_existente}"
                    else:
                        # Faz upload e transcrição se o arquivo não tiver transcrição prévia
                        audio_url = upload_audio(file_path)
                        print(file_path)
                        if audio_url:
                            transcription = transcrever_audio_assemblyai(audio_url)
                            if transcription:
                                row_cells[2].text = f"ÁUDIO\nTranscrição: '{transcription}'"
                                salvar_transcricao(hash_arquivo, transcription, from_field, to_field, timestamp)
                            else:
                                row_cells[2].text = "Falha ao transcrever o áudio."
                        else:
                            row_cells[2].text = "Falha ao enviar o áudio."
                else:
                    # Mensagem se o arquivo de áudio não for encontrado
                    row_cells[2].text = "Áudio deletado"
            
            elif attachment.lower().endswith(('.jpg', '.jpeg', '.png', '.gif')):
                row_cells[2].text = f"IMAGEM: {attachment}"
                if os.path.exists(file_path):
                    try:
                        paragraph = row_cells[2].paragraphs[0]
                        run = paragraph.add_run()
                        run.add_picture(file_path, width=Cm(7))
                    except Exception as e:
                        row_cells[2].text = f"Erro ao inserir imagem: {e}"
                else:
                    row_cells[2].text = "Imagem não encontrada no caminho especificado."
            else:
                row_cells[2].text = str(body) if body else "Sem conteúdo"  # Garantia de que é string
                #print(body)

        # Atualiza a barra de progresso a cada linha
        linhas_processadas += 1
        progresso = linhas_processadas / total_linhas
        progress_bar.progress(progresso)

        row_cells[3].text = str(timestamp)

    lote_file_path = f"lote_{lote_num}.docx"
    doc.save(lote_file_path)
    return lote_file_path


#
#####
######## VERIFICAR ESSA FUNÇÃO NÃO ESTÁ COLOCANDO A IMAGEM
def criar_documento_para_lote2(df_lote, audio_dir, lote_num, progress_bar, linhas_processadas, total_linhas):
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'From'
    hdr_cells[2].text = 'Body'
    hdr_cells[3].text = 'Timestamp-Time'
    item = 0
           
    for index, row in df_lote.iterrows():
        item += 1
        from_field = row['From']
        to_field = row.get('To', '')
        timestamp = row.get('Timestamp-Time', row.get('Timestamp: Time', ''))
        body = row.get('Body', '')
        attachment = row.get('Attachment #1', None)
        label = row.get('Label', '')
        #is_forwarded = row.get('Label', '') == 'Forwarded' or 'Forwarded '
        is_forwarded = "Forwarded" in label.strip()
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(item)
        row_cells[1].text = from_field

        # Prefixo "ENCAMINHADO" se for uma mensagem encaminhada
        prefixo_encaminhado = "ENCAMINHADO\n" if is_forwarded else ""

        if pd.notna(attachment):
            file_path = os.path.join(audio_dir, attachment)
            
            if attachment.endswith('.opus'):
                if os.path.exists(file_path):
                    hash_arquivo = gerar_hash_arquivo(file_path)
                    transcricao_existente = buscar_transcricao(hash_arquivo)
                    
                    if transcricao_existente:
                        row_cells[2].text = f"{prefixo_encaminhado}ÁUDIO\nTranscrição: {transcricao_existente}"
                    else:
                        audio_url = upload_audio(file_path)
                        if audio_url:
                            transcription = transcrever_audio_assemblyai(audio_url)
                            if transcription:
                                row_cells[2].text = f"{prefixo_encaminhado}ÁUDIO\nTranscrição: '{transcription}'"
                                salvar_transcricao(hash_arquivo, transcription, from_field, to_field, timestamp)
                            else:
                                row_cells[2].text = f"{prefixo_encaminhado}Falha ao transcrever o áudio."
                        else:
                            row_cells[2].text = f"{prefixo_encaminhado}Falha ao enviar o áudio."
                else:
                    row_cells[2].text = f"{prefixo_encaminhado}Áudio deletado"
            
            elif attachment.lower().endswith(('.jpg', '.jpeg', '.png', '.gif')):
                row_cells[2].text = f"{prefixo_encaminhado}IMAGEM: {attachment}"
                if os.path.exists(file_path):
                    try:
                        paragraph = row_cells[2].paragraphs[0]
                        run = paragraph.add_run()
                        run.add_picture(file_path, width=Cm(7))
                    except Exception as e:
                        row_cells[2].text = f"{prefixo_encaminhado}Erro ao inserir imagem: {e}"
                else:
                    row_cells[2].text = f"{prefixo_encaminhado}Imagem não encontrada no caminho especificado."
            else:
                row_cells[2].text = f"{prefixo_encaminhado}{str(body) if body else 'Sem conteúdo'}"

        else:
            row_cells[2].text = f"{prefixo_encaminhado}{str(body) if body else 'Sem conteúdo'}"

        linhas_processadas += 1
        progresso = linhas_processadas / total_linhas
        progress_bar.progress(progresso)

        row_cells[3].text = str(timestamp)

    lote_file_path = f"lote_{lote_num}.docx"
    doc.save(lote_file_path)
    return lote_file_path



def processar_em_lotes(file, audio_dir, progress_bar, filtro_tag):
    df = pd.read_excel(file, engine='openpyxl', header=1).dropna(how='all')

    # Aplicar o filtro na coluna "Tag" se necessário
    if filtro_tag:
        df = df[df['Tag'].notna()]  # Filtra apenas linhas onde "Tag" não está vazia
    #print(filtro_tag)

    df = df.fillna('')

    total_linhas = len(df)  # Total de linhas da planilha
    
    linhas_processadas = 0  # Contador de linhas processadas
    lotes = [df[i:i + 500] for i in range(0, len(df), 500)]
    #total_passos = len(lotes)

    documentos_gerados = []
    #progresso = 0

    for lote_num, df_lote in enumerate(lotes, start=1):
        lote_file = criar_documento_para_lote2(df_lote, audio_dir, lote_num, progress_bar, linhas_processadas, total_linhas)
        documentos_gerados.append(lote_file)
        #progresso += 1

        linhas_processadas += len(df_lote)
        
    return documentos_gerados

# Função para recriar o conteúdo de cada lote no documento final, incluindo as imagens
def recriar_documento_final(documentos, audio_dir, doc_final_path="resultado_transcricoes_final.docx"):
    doc_final = Document()

    for idx, doc_path in enumerate(documentos):
        doc_lote = Document(doc_path)

        for table in doc_lote.tables:
            new_table = doc_final.add_table(rows=1, cols=4)
            new_table.autofit = False

            # Recriar cabeçalho
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Item'
            hdr_cells[1].text = 'From'
            hdr_cells[2].text = 'Body'
            hdr_cells[3].text = 'Timestamp-Time'
            item = 0
            

            for row in table.rows[1:]:  # Ignorar cabeçalho
                item = item + 1
                new_row = new_table.add_row().cells
                for i, cell in enumerate(row.cells):
                    if i == 2 and cell.text.startswith("IMAGEM:"):
                        # Extrair o nome do arquivo de imagem
                        attachment_name = cell.text.replace("IMAGEM:", "").strip()
                        image_path = os.path.join(audio_dir, attachment_name)

                        # Inserir a imagem se o caminho existir
                        if os.path.exists(image_path):
                            try:
                                paragraph = new_row[i].paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(image_path, width=Cm(7))
                            except Exception as e:
                                new_row[i].text = f"Erro ao inserir imagem: {e}"
                        else:
                            new_row[i].text = f"Imagem '{attachment_name}' não encontrada no caminho especificado."
                    else:
                        new_row[i].text = cell.text

            # Adicionar quebra de página entre lotes
            if idx < len(documentos) - 1:
                doc_final.add_page_break()


    doc_final.save(doc_final_path)

    formatar_tabela_documento(doc_final_path)

    return doc_final_path

def formatar_tabela_documento(doc_path):
    doc = Document(doc_path)
    
    # Configuração de margens mínimas na página
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1)
        section.right_margin = Cm(0.7)
    
    # Configuração de estilo das colunas
    largura_coluna_item = Cm(1.1)
    largura_coluna_from = Cm(4)
    largura_coluna_body = Cm(12)
    largura_coluna_timestamp = Cm(3.6)

    # Dicionário para mapear conteúdos únicos para cores
    cores_conteudos = {
        "System Message System Message": "D9D9D9"  # Cor fixa para mensagens de sistema
    }
    cores_alternadas = ["B6DDE8", "D6E3BC"]  # Cores para conteúdos alternados
    cor_indice = 0  # Índice para alternar entre as cores

    for table in doc.tables:
        # Formatação da primeira linha como cabeçalho
        header_row = table.rows[0]
        header_row.cells[0].text = "Item"
        header_row.cells[1].text = "De"
        header_row.cells[2].text = "Mensagem"
        header_row.cells[3].text = "Data e Hora"
        for cell in header_row.cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Aplicação de largura de coluna, formatação de linhas e bordas
        for row_idx, row in enumerate(table.rows):

            # Aplicação da largura nas três primeiras células da linha
            ajustar_largura_celula(row.cells[0], largura_coluna_item)
            ajustar_largura_celula(row.cells[1], largura_coluna_from)
            ajustar_largura_celula(row.cells[2], largura_coluna_body)
            ajustar_largura_celula(row.cells[3], largura_coluna_timestamp)

            if row_idx == 0:
                # Configuração do cabeçalho
                #row.cells[0].width = largura_coluna_from
                #row.cells[1].width = largura_coluna_body
                #row.cells[2].width = largura_coluna_timestamp

                # Aplicar bordas ao cabeçalho
                for cell in row.cells:
                    definir_borda_celula(cell, "single", "000000")  # Cor da borda preta

            else:
                from_text = row.cells[1].text.strip()

                # Definir cor para o conteúdo da linha com base no "From"
                if from_text in cores_conteudos:
                    cor_linha = cores_conteudos[from_text]
                else:
                    cor_linha = cores_alternadas[cor_indice % len(cores_alternadas)]
                    cores_conteudos[from_text] = cor_linha
                    cor_indice += 1

                # Aplicação da cor de fundo e bordas para cada célula da linha
                for cell in row.cells:
                    # Aplicar cor de fundo
                    cell_tcPr = cell._element.find(qn("w:tcPr"))
                    if cell_tcPr is None:
                        cell_tcPr = OxmlElement("w:tcPr")
                        cell._element.append(cell_tcPr)
                    
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:fill"), cor_linha)
                    cell_tcPr.append(shading_elm)

                    # Aplicar bordas em cada célula
                    definir_borda_celula(cell, "single", "000000")  # Cor da borda preta
                    
                # Centralizar conteúdo da coluna "De"
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Salvar o documento com a formatação aplicada
    doc.save(doc_path)

def definir_borda_celula(cell, borda_tipo="single", borda_cor="000000"):
    # Verificar e adicionar `w:tcPr` se não existir
    tcPr = cell._element.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        cell._element.append(tcPr)
    
    for borda in ["top", "bottom", "left", "right"]:
        borda_element = OxmlElement(f"w:{borda}")
        borda_element.set(qn("w:val"), borda_tipo)
        borda_element.set(qn("w:sz"), "4")  # Espessura da borda
        borda_element.set(qn("w:color"), borda_cor)
        tcPr.append(borda_element)

def ajustar_largura_celula(cell, largura):
    # Adiciona manualmente o elemento `w:tcPr` se ele não existir
    tcPr = cell._element.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        cell._element.append(tcPr)
    
    # Adiciona ou atualiza o elemento de largura da célula
    cell_width = OxmlElement("w:tcW")
    cell_width.set(qn("w:w"), str(int(largura.cm * 567)))  # Converter Cm para unidade de medida interna do Word
    cell_width.set(qn("w:type"), "dxa")
    tcPr.append(cell_width)
    

def anonimizar_interlocutores4(docx_path, output_path):
    st.write("INICIOU ANONIMIZAÇÃO")
    
    # Carregar o documento .docx
    doc = Document(docx_path)
    
    # Etapa 1: Mapeamento dos identificadores principais para Interlocutor 1 e Interlocutor 2
    interlocutor_map = {}
    interlocutor_count = 1  # Contador para nomeação dos interlocutores
    
    # Primeiro, percorremos todas as tabelas para identificar os interlocutores
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            # Ignorar a primeira linha (títulos da tabela)
            if i == 0:
                continue

            # Verificar se a linha tem pelo menos duas colunas
            if len(row.cells) >= 2:
                cell_value_2 = row.cells[1].text.strip()
                
                # Ignorar linhas da segunda coluna com "System Message"
                if cell_value_2 != "System Message" and cell_value_2 != "System Message System Message":
                    # Extrair o identificador principal (exemplo: número de telefone com domínio)
                    identifier_match = re.search(r"\b\d{10,}@\w+\.\w+\b", cell_value_2)
                    if identifier_match:
                        identifier = identifier_match.group(0)
                        
                        # Mapear o identificador se ainda não estiver mapeado e temos menos de 2 interlocutores
                        if identifier not in interlocutor_map and interlocutor_count <= 2:
                            interlocutor_map[identifier] = f"Interlocutor {interlocutor_count}"
                            interlocutor_count += 1
    
    # Exibir o mapeamento dos interlocutores
    st.write("Mapeamento dos Interlocutores:")
    for original, anonimo in interlocutor_map.items():
        st.write(f"{anonimo}: {original}")
    
    # Etapa 2: Substituição dos identificadores no documento
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            # Ignorar a primeira linha (títulos da tabela)
            if i == 0:
                continue
            
            # Verificar se a linha tem pelo menos duas colunas
            if len(row.cells) >= 2:
                # Substituição total na segunda coluna, se não for "System Message"
                cell_value_2 = row.cells[1].text.strip()
                if cell_value_2 != "System Message" and cell_value_2 != "System Message System Message":
                    for identifier, anonimo in interlocutor_map.items():
                        if identifier in cell_value_2:
                            # Redefinir o texto da célula completamente para o nome do interlocutor
                            row.cells[1].text = anonimo  # Substituição total e direta
                
                # Verificação para ignorar células com imagens na terceira coluna
                if len(row.cells) >= 3:
                    cell = row.cells[2]
                    has_image = any([shape for shape in cell._element.xpath('.//w:drawing')])
                    
                    if not has_image:  # Apenas modificar se não houver imagem
                        for paragraph in cell.paragraphs:
                            cell_text = paragraph.text
                            # Realizar substituição no texto do parágrafo com regex para capturar trechos extras
                            for identifier, anonimo in interlocutor_map.items():
                                # Substituir qualquer ocorrência do identificador seguido por caracteres adicionais
                                cell_text = re.sub(rf"\b{re.escape(identifier)}\b(?:\.\w+)?", anonimo, cell_text)
                            # Atualizar o texto do parágrafo
                            paragraph.text = cell_text
    
    # Salvar o documento modificado
    doc.save(output_path)
    st.write("Anonimização concluída. Documento salvo em:", output_path)


def anonimizar_interlocutores5(docx_path, output_path):
    st.write("INICIOU ANONIMIZAÇÃO")
    
    # Carregar o documento .docx
    doc = Document(docx_path)
    
    # Etapa 1: Mapeamento detalhado dos identificadores
    interlocutor_map = {}
    interlocutor_details = {}
    interlocutor_count = 1  # Contador para nomeação dos interlocutores
    
    # Primeiro, percorremos todas as tabelas para identificar os interlocutores
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            # Ignorar a primeira linha (títulos da tabela)
            if i == 0:
                continue

            # Verificar se a linha tem pelo menos duas colunas
            if len(row.cells) >= 2:
                cell_value_2 = row.cells[1].text.strip()
                
                # Ignorar linhas da segunda coluna com "System Message"
                if cell_value_2 not in ["System Message", "System Message System Message"]:
                    # Extrair o identificador principal e o nome, permitindo múltiplos espaços
                    match = re.match(r"(\d{10,}@[\w.]+\.\w+)\s+(.*)", cell_value_2)
                    if match:
                        number = match.group(1)  # O número de telefone com domínio
                        name = match.group(2)    # O nome do interlocutor
                        full_identifier = f"{number} {name}"
                        
                        # Mapear o identificador se ainda não estiver mapeado e temos menos de 2 interlocutores
                        if full_identifier not in interlocutor_map and interlocutor_count <= 2:
                            interlocutor_map[full_identifier] = f"Interlocutor {interlocutor_count}"
                            interlocutor_details[full_identifier] = {
                                "name": name,
                                "number": number,
                                "anon_name": f"Nome {interlocutor_count}",
                                "anon_number": f"Número {interlocutor_count}"
                            }
                            interlocutor_count += 1
    
    # Exibir o mapeamento detalhado dos interlocutores
    st.write("Mapeamento dos Interlocutores:")
    for original, anonimo in interlocutor_map.items():
        details = interlocutor_details[original]
        st.write(f"{anonimo} = {original}")
        st.write(f"  {details['anon_name']}: {details['name']}")
        st.write(f"  {details['anon_number']}: {details['number']}")
    
    # Etapa 2: Substituição dos identificadores no documento
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            # Ignorar a primeira linha (títulos da tabela)
            if i == 0:
                continue
            
            # Verificar se a linha tem pelo menos duas colunas
            if len(row.cells) >= 2:
                cell_value_2 = row.cells[1].text.strip()
                
                # Substituição total na segunda coluna, se não for "System Message"
                if cell_value_2 not in ["System Message", "System Message System Message"]:
                    for full_identifier, anonimo in interlocutor_map.items():
                        if full_identifier in cell_value_2:
                            # Redefinir o texto da célula completamente para o nome do interlocutor
                            row.cells[1].text = anonimo  # Substituição total e direta
                
                # Caso seja "System Message" ou "System Message System Message", verificar a coluna 3
                elif cell_value_2 in ["System Message", "System Message System Message"]:
                    if len(row.cells) >= 3:
                        cell = row.cells[2]
                        has_image = any([shape for shape in cell._element.xpath('.//w:drawing')])
                        
                        if not has_image:  # Apenas modificar se não houver imagem
                            for paragraph in cell.paragraphs:
                                cell_text = paragraph.text
                                # Realizar substituição com regex para o nome e o número dos interlocutores
                                for full_identifier, details in interlocutor_details.items():
                                    cell_text = re.sub(rf"\b{re.escape(details['number'])}\b(?:\.\w+)?", details["anon_number"], cell_text)
                                    cell_text = re.sub(rf"\b{re.escape(details['name'])}\b", details["anon_name"], cell_text)
                                # Atualizar o texto do parágrafo
                                paragraph.text = cell_text
    
    # Salvar o documento modificado
    doc.save(output_path)
    st.write("Anonimização concluída. Documento salvo em:", output_path)
