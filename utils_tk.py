import os
import pandas as pd
import time
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt
from assemblyai import upload_audio, transcrever_audio_assemblyai
from PIL import UnidentifiedImageError, Image
from db import salvar_transcricao, buscar_transcricao
import streamlit as st
from docx.enum.text import WD_ALIGN_PARAGRAPH
import hashlib
import re
from PIL import Image, ImageOps
import cv2

from tempfile import NamedTemporaryFile

def redimensionar_imagem(image_path, max_width=7):
    """Redimensiona a imagem para ter uma largura máxima de 7 cm, mantendo a proporção."""
    with Image.open(image_path) as img:
        width, height = img.size
        aspect_ratio = height / width
        new_width = min(width, max_width * 37.7953)  # Conversão para pixels (1 cm ≈ 37.7953 pixels)
        new_height = int(new_width * aspect_ratio)
        img = img.resize((int(nesw_width), new_height), Image.LANCZOS)
        # Salvar a imagem temporária redimensionada
        temp_path = f"{image_path}_temp_resized.jpg"
        img.save(temp_path, format="JPEG", quality=85)
        return temp_path

# Função para adicionar bordas a todas as células da tabela
def adicionar_bordas_a_tabela(table):
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')  # Define as bordas como simples
                border.set(qn('w:sz'), '4')        # Espessura da borda
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')  # Cor preta
                tc_borders.append(border)
            
            tc_pr.append(tc_borders)

def verificar_arquivos_na_pasta(file, audio_dir):
    # Carregar o arquivo Excel
    df = pd.read_excel(file, engine='openpyxl', header=1)

    # Filtrar linhas em branco
    df = df.dropna(how='all')

    # Iterar pelas linhas e verificar se pelo menos um arquivo existe na pasta
    for index, row in df.iterrows():
        attachment = row.get('Attachment #1', row.get('Anexo #1', None))
        if pd.notna(attachment):
            file_path = os.path.join(audio_dir, attachment)
            if os.path.exists(file_path):
                return True  # Arquivo encontrado, a pasta está correta
        else:
            return True #retorna True poruqe não tem nenhum arquivo anexo na planilha
    return False  # Nenhum arquivo encontrado, a pasta está incorreta

def gerar_hash_arquivo(caminho_arquivo):
    hash_sha256 = hashlib.sha256()
    with open(caminho_arquivo, "rb") as f:
        # Lê o arquivo em blocos para calcular o hash (eficiente para arquivos grandes)
        for bloco in iter(lambda: f.read(4096), b""):
            hash_sha256.update(bloco)
    return hash_sha256.hexdigest()

def criar_documento_para_lote(df_lote, audio_dir, lote_num, progress_bar, linhas_processadas, total_linhas):
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'From'
    hdr_cells[2].text = 'Body'
    hdr_cells[3].text = 'Timestamp-Time'
    item = 0
           
    for index, row in df_lote.iterrows():
        item += 1
        from_field = row.get('From', row.get('De', ''))
        to_field = row.get('To', row.get('Para', ''))
        timestamp = row.get('Timestamp-Time', row.get('Timestamp: Time', row.get('Marcação de tempo-Hora', '')))
        body = row.get('Body', row.get('Corpo', ''))
        attachment = row.get('Attachment #1', row.get('Anexo #1', None))
        label = row.get('Label', row.get('Rótulo', ''))
        deleted = row.get('Deleted', row.get('Excluído', row.get('Deleted - Instant Message', '')))
        location = row.get('Location', '')  # Verifica a coluna "Location"

        # Verifica se a mensagem foi encaminhada
        is_forwarded = "Forwarded" in label.strip() or "Encaminhado" in label.strip()
        
        # Verifica se a mensagem foi marcada como excluída
        is_deleted = isinstance(deleted, str) and deleted.strip()

        # Verifica se existe uma localização
        has_location = isinstance(location, str) and location.strip()
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(item)
        row_cells[1].text = from_field

        # Inicializa o conteúdo de "Body" com os rótulos apropriados
        message_label = ""
        if is_forwarded:
            message_label += "➡️ Encaminhado\n"
        if is_deleted:
            message_label += "❌ Mensagem Excluída\n"
        if has_location:
            message_label += f"📍 Localização: {location.strip()}\n"  # Adiciona o rótulo de Localização
            #print("\nattachment attachment attachment: ",attachment)
            #print("ITEM: ",item)
        # Verifica o tipo de conteúdo para "Body" ou anexo
        if pd.notna(attachment):
            file_path = os.path.join(audio_dir, attachment)
            
            if attachment.endswith('.opus'):
                row_cells[2].text = message_label
                if os.path.exists(file_path):
                    hash_arquivo = gerar_hash_arquivo(file_path)
                    transcricao_existente = buscar_transcricao(hash_arquivo)
                    
                    if transcricao_existente:
                        row_cells[2].text += f"🔉 Áudio\nTranscrição: {transcricao_existente}"
                    else:
                        audio_url = upload_audio(file_path)
                        if audio_url:
                            transcription = transcrever_audio_assemblyai(audio_url)
                            if transcription:
                                row_cells[2].text += f"🔉 Áudio\nTranscrição: '{transcription}'"
                                salvar_transcricao(hash_arquivo, transcription, from_field, to_field, timestamp)
                            else:
                                row_cells[2].text += "Falha ao transcrever o áudio."
                        else:
                            row_cells[2].text += "Falha ao enviar o áudio."
                else:
                    row_cells[2].text += "Áudio deletado"
            
            elif attachment.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.webp', '.heic')):
                row_cells[2].text = f"{message_label}📸 Imagem: {attachment}"
            
            elif attachment.lower().endswith(('.pdf', '.docx', '.xlsx')):
                row_cells[2].text = f"{message_label}📁 Arquivo: \"{attachment}\""
                
            elif attachment.lower().endswith(('.mp4')):
                row_cells[2].text = f"{message_label}📽️ Vídeo: {attachment}"
                
            elif attachment.lower().endswith(('.thumb')):
                row_cells[2].text = f"{message_label}📁 Arquivo: {attachment}"
            elif "Shared" in attachment.strip():
                attachment2 = re.sub(r'(_x000d_|[\r\n])', '\n', attachment).strip()
                row_cells[2].text = f"{message_label}🪪 Contato: {attachment2}"
            else:
                row_cells[2].text = f"{message_label}{str(body) if body else 'Arquivo não recuperado!'}"
                
        else:
            # Conteúdo genérico sem anexo
            row_cells[2].text = f"{message_label}{str(body) if body else 'Sem conteúdo'}"

        # Atualiza a barra de progresso e define o timestamp
        linhas_processadas += 1
        #progresso = linhas_processadas / total_linhas
        #progress_bar.progress(progresso)
        progress = (linhas_processadas / total_linhas) * 100
        progress_bar(progress)  # Atualiza a barra de progresso com o valor atual


        row_cells[3].text = str(timestamp)

    # Salva o arquivo do lote
    lote_file_path = f"lote_{lote_num}.docx"
    doc.save(lote_file_path)
    return lote_file_path

def processar_em_lotes(file, audio_dir, progress_bar, filtro_tag):
    df = pd.read_excel(file, engine='openpyxl', header=1).dropna(how='all')

    # Aplicar o filtro na coluna "Tag" se necessário
    if filtro_tag:
        if df['Tag']:
            df = df[df['Tag'].notna()]
        elif df['Tag Note - Instant Message']:
            df = df[df['Tag Note - Instant Message'].notna()]  # Filtra apenas linhas onde "Tag" não está vazia
        else:
            df = df[df['Nota da etiqueta'].notna()]
    #print(filtro_tag)

    df = df.fillna('')

    total_linhas = len(df)  # Total de linhas da planilha
    
    linhas_processadas = 0  # Contador de linhas processadas
    lotes = [df[i:i + 500] for i in range(0, len(df), 500)]
    #total_passos = len(lotes)

    documentos_gerados = []
    #progresso = 0

    for lote_num, df_lote in enumerate(lotes, start=1):
        lote_file = criar_documento_para_lote(df_lote, audio_dir, lote_num, progress_bar, linhas_processadas, total_linhas)
        documentos_gerados.append(lote_file)
        #progresso += 1

        linhas_processadas += len(df_lote)
        
    return documentos_gerados


def capturar_frame(video_path, output_path, frame_number=10):
    """Captura um frame específico de um vídeo e salva como uma imagem .jpeg."""
    try:
        vidcap = cv2.VideoCapture(video_path)
        vidcap.set(cv2.CAP_PROP_POS_FRAMES, frame_number)  # Define o frame específico
        success, frame = vidcap.read()
        if success:
            cv2.imwrite(output_path, frame)  # Salva o frame específico como imagem
        vidcap.release()
        return success
    except Exception as e:
        print(f"Erro ao capturar o frame {frame_number} do vídeo {video_path}: {e}")
        return False

def recriar_documento_final(documentos, audio_dir, doc_final_path="resultado_transcricoes_final.docx", frame_number=10):
    doc_final = Document()

    for idx, doc_path in enumerate(documentos):
        doc_lote = Document(doc_path)

        for table in doc_lote.tables:
            new_table = doc_final.add_table(rows=1, cols=4)
            new_table.autofit = False

            # Recriar cabeçalho
            hdr_cells = new_table.rows[0].cells
            hdr_cells[0].text = 'Item'
            hdr_cells[1].text = 'From'
            hdr_cells[2].text = 'Body'
            hdr_cells[3].text = 'Timestamp-Time'
            item = 0

            for row in table.rows[1:]:  # Ignorar cabeçalho
                item += 1
                new_row = new_table.add_row().cells
                for i, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()

                    # Verifica se a célula contém uma imagem ou vídeo e extrai o nome do arquivo
                    if i == 2 and ("📸 Imagem:" in cell_text or "📽️ Vídeo:" in cell_text):
                        forwarded = "➡️ Encaminhado" in cell_text
                        deleted = "❌ Mensagem Excluída" in cell_text

                        # Remove rótulos antes de buscar o nome do anexo
                        clean_text = cell_text.replace("➡️ Encaminhado\n", "").replace("❌ Mensagem Excluída\n", "").strip()
                        attachment_name = clean_text.replace("📸 Imagem:", "").replace("📽️ Vídeo:", "").strip()
                        file_path = os.path.join(audio_dir, attachment_name)

                        # Configura o rótulo inicial para o texto
                        label_text = ""
                        if forwarded:
                            label_text += "➡️ Encaminhada\n"
                        if deleted:
                            label_text += "❌ Mensagem Excluída\n"

                        # Verifica a extensão do arquivo
                        _, ext = os.path.splitext(attachment_name)
                        ext = ext.lower()

                        # Tratamento específico para vídeos .mp4
                        if ext == '.mp4':
                            # Adiciona "📽️ Vídeo" seguido do nome do arquivo
                            label_text += f"📽️ Vídeo: {attachment_name}\n"
                            new_row[i].text = label_text

                            # Captura o frame especificado do vídeo
                            temp_image_path = file_path.replace(ext, "_frame.jpeg")
                            if capturar_frame(file_path, temp_image_path, frame_number=frame_number):
                                # Insere o frame capturado
                                if os.path.exists(temp_image_path):
                                    try:
                                        paragraph = new_row[i].paragraphs[0]
                                        run = paragraph.add_run()
                                        run.add_picture(temp_image_path, width=Cm(5))  # Tamanho menor para prévia do vídeo
                                    except Exception as e:
                                        new_row[i].text += f"Erro ao inserir prévia do vídeo: {e}"
                                    finally:
                                        os.remove(temp_image_path)  # Remove a imagem temporária
                            else:
                                new_row[i].text += "Erro ao capturar o frame especificado do vídeo."

                        # Tratamento para outras imagens
                        elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.heic']:
                            if ext == '.webp':
                                label_text += f"Sticker: {attachment_name}\n"  # Adiciona "Sticker" seguido do nome do arquivo
                                image_size = Cm(3)  # Tamanho menor para stickers
                            else:
                                label_text += f"📸 Imagem: {attachment_name}\n"  # Adiciona "📸 Imagem" seguido do nome do arquivo
                                image_size = Cm(7)  # Tamanho padrão para outras imagens
                            new_row[i].text = label_text

                            # Converte .webp e .heic para .jpeg, se necessário
                            if ext in ['.webp', '.heic']:
                                try:
                                    with Image.open(file_path) as img:
                                        img = img.convert("RGB")
                                        temp_image_path = file_path.replace(ext, ".jpeg")
                                        img.save(temp_image_path, format="JPEG")
                                    file_path = temp_image_path
                                except Exception as e:
                                    new_row[i].text += f"Erro ao processar imagem {ext}: {e}"
                                    continue

                            # Insere a imagem com o tamanho especificado
                            if os.path.exists(file_path):
                                try:
                                    paragraph = new_row[i].paragraphs[0]
                                    run = paragraph.add_run()
                                    run.add_picture(file_path, width=image_size)
                                except Exception as e:
                                    new_row[i].text += f"Erro ao inserir imagem: {e}"
                            else:
                                new_row[i].text += f"Imagem '{attachment_name}' não encontrada no caminho especificado."

                            # Remove o arquivo temporário se tiver sido criado
                            if ext in ['.webp', '.heic'] and os.path.exists(file_path):
                                os.remove(file_path)

                    else:
                        # Para outras células, transfere o texto como está
                        new_row[i].text = cell.text

            # Adicionar quebra de página entre lotes
            if idx < len(documentos) - 1:
                doc_final.add_page_break()

    # Salva o documento final
    doc_final.save(doc_final_path)
    formatar_tabela_documento(doc_final_path)

    return doc_final_path


def formatar_tabela_documento(doc_path):
    doc = Document(doc_path)
    
    # Configuração de margens mínimas na página
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1)
        section.right_margin = Cm(0.7)
    
    # Configuração de estilo das colunas
    largura_coluna_item = Cm(1.1)
    largura_coluna_from = Cm(4)
    largura_coluna_body = Cm(12)
    largura_coluna_timestamp = Cm(3.6)

    # Dicionário para mapear conteúdos únicos para cores
    cores_conteudos = {
        "System Message System Message": "D9D9D9"  # Cor fixa para mensagens de sistema
    }
    cores_alternadas = ["B6DDE8", "D6E3BC"]  # Cores para conteúdos alternados
    cor_indice = 0  # Índice para alternar entre as cores

    for table in doc.tables:
        # Formatação da primeira linha como cabeçalho
        header_row = table.rows[0]
        header_row.cells[0].text = "Item"
        header_row.cells[1].text = "De"
        header_row.cells[2].text = "Mensagem"
        header_row.cells[3].text = "Data e Hora"
        for cell in header_row.cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Aplicação de largura de coluna, formatação de linhas e bordas
        for row_idx, row in enumerate(table.rows):

            # Aplicação da largura nas três primeiras células da linha
            ajustar_largura_celula(row.cells[0], largura_coluna_item)
            ajustar_largura_celula(row.cells[1], largura_coluna_from)
            ajustar_largura_celula(row.cells[2], largura_coluna_body)
            ajustar_largura_celula(row.cells[3], largura_coluna_timestamp)

            if row_idx == 0:
                # Configuração do cabeçalho
                #row.cells[0].width = largura_coluna_from
                #row.cells[1].width = largura_coluna_body
                #row.cells[2].width = largura_coluna_timestamp

                # Aplicar bordas ao cabeçalho
                for cell in row.cells:
                    definir_borda_celula(cell, "single", "000000")  # Cor da borda preta

            else:
                from_text = row.cells[1].text.strip()

                # Definir cor para o conteúdo da linha com base no "From"
                if from_text in cores_conteudos:
                    cor_linha = cores_conteudos[from_text]
                else:
                    cor_linha = cores_alternadas[cor_indice % len(cores_alternadas)]
                    cores_conteudos[from_text] = cor_linha
                    cor_indice += 1

                # Aplicação da cor de fundo e bordas para cada célula da linha
                for cell in row.cells:
                    # Aplicar cor de fundo
                    cell_tcPr = cell._element.find(qn("w:tcPr"))
                    if cell_tcPr is None:
                        cell_tcPr = OxmlElement("w:tcPr")
                        cell._element.append(cell_tcPr)
                    
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:fill"), cor_linha)
                    cell_tcPr.append(shading_elm)

                    # Aplicar bordas em cada célula
                    definir_borda_celula(cell, "single", "000000")  # Cor da borda preta
                    
                # Centralizar conteúdo da coluna "De"
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Salvar o documento com a formatação aplicada
    doc.save(doc_path)

def definir_borda_celula(cell, borda_tipo="single", borda_cor="000000"):
    # Verificar e adicionar `w:tcPr` se não existir
    tcPr = cell._element.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        cell._element.append(tcPr)
    
    for borda in ["top", "bottom", "left", "right"]:
        borda_element = OxmlElement(f"w:{borda}")
        borda_element.set(qn("w:val"), borda_tipo)
        borda_element.set(qn("w:sz"), "4")  # Espessura da borda
        borda_element.set(qn("w:color"), borda_cor)
        tcPr.append(borda_element)

def ajustar_largura_celula(cell, largura):
    # Adiciona manualmente o elemento `w:tcPr` se ele não existir
    tcPr = cell._element.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        cell._element.append(tcPr)
    
    # Adiciona ou atualiza o elemento de largura da célula
    cell_width = OxmlElement("w:tcW")
    cell_width.set(qn("w:w"), str(int(largura.cm * 567)))  # Converter Cm para unidade de medida interna do Word
    cell_width.set(qn("w:type"), "dxa")
    tcPr.append(cell_width)
    

def anonimizar_interlocutores(docx_path, output_path):
    st.write("INICIOU ANONIMIZAÇÃO")
    
    # Carregar o documento .docx
    doc = Document(docx_path)
    
    # Etapa 1: Mapeamento detalhado dos identificadores
    interlocutor_map = {}
    interlocutor_details = {}
    interlocutor_count = 1  # Contador para nomeação dos interlocutores
    
    # Primeiro, percorremos todas as tabelas para identificar os interlocutores
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            # Ignorar a primeira linha (títulos da tabela)
            if i == 0:
                continue

            # Verificar se a linha tem pelo menos duas colunas
            if len(row.cells) >= 2:
                cell_value_2 = row.cells[1].text.strip()
                
                # Ignorar linhas da segunda coluna com "System Message"
                if cell_value_2 not in ["System Message", "System Message System Message"]:
                    # Extrair o identificador principal e o nome, permitindo múltiplos espaços
                    match = re.match(r"(\d{10,}@[\w.]+\.\w+)\s+(.*)", cell_value_2)
                    if match:
                        number = match.group(1)  # O número de telefone com domínio
                        name = match.group(2)    # O nome do interlocutor
                        full_identifier = f"{number} {name}"
                        
                        # Mapear o identificador se ainda não estiver mapeado e temos menos de 2 interlocutores
                        if full_identifier not in interlocutor_map and interlocutor_count <= 2:
                            interlocutor_map[full_identifier] = f"Interlocutor {interlocutor_count}"
                            interlocutor_details[full_identifier] = {
                                "name": name,
                                "number": number,
                                "anon_name": f"Nome {interlocutor_count}",
                                "anon_number": f"Número {interlocutor_count}"
                            }
                            interlocutor_count += 1
    
    # Exibir o mapeamento detalhado dos interlocutores
    st.write("Mapeamento dos Interlocutores:")
    for original, anonimo in interlocutor_map.items():
        details = interlocutor_details[original]
        st.write(f"{anonimo} = {original}")
        st.write(f"  {details['anon_name']}: {details['name']}")
        st.write(f"  {details['anon_number']}: {details['number']}")
    
    # Etapa 2: Substituição dos identificadores no documento
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            # Ignorar a primeira linha (títulos da tabela)
            if i == 0:
                continue
            
            # Verificar se a linha tem pelo menos duas colunas
            if len(row.cells) >= 2:
                cell_value_2 = row.cells[1].text.strip()
                
                # Substituição total na segunda coluna, se não for "System Message"
                if cell_value_2 not in ["System Message", "System Message System Message"]:
                    for full_identifier, anonimo in interlocutor_map.items():
                        if full_identifier in cell_value_2:
                            # Redefinir o texto da célula completamente para o nome do interlocutor
                            row.cells[1].text = anonimo  # Substituição total e direta
                
                # Caso seja "System Message" ou "System Message System Message", verificar a coluna 3
                elif cell_value_2 in ["System Message", "System Message System Message"]:
                    if len(row.cells) >= 3:
                        cell = row.cells[2]
                        has_image = any([shape for shape in cell._element.xpath('.//w:drawing')])
                        
                        if not has_image:  # Apenas modificar se não houver imagem
                            for paragraph in cell.paragraphs:
                                cell_text = paragraph.text
                                # Realizar substituição com regex para o nome e o número dos interlocutores
                                for full_identifier, details in interlocutor_details.items():
                                    cell_text = re.sub(rf"\b{re.escape(details['number'])}\b(?:\.\w+)?", details["anon_number"], cell_text)
                                    cell_text = re.sub(rf"\b{re.escape(details['name'])}\b", details["anon_name"], cell_text)
                                # Atualizar o texto do parágrafo
                                paragraph.text = cell_text
    
    # Salvar o documento modificado
    doc.save(output_path)
    st.write("Anonimização concluída. Documento salvo em:", output_path)
